import io
from copy import deepcopy
from pathlib import Path

import streamlit as st
from docx import Document

APP_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = APP_DIR / "template.docx"

st.set_page_config(page_title="Circuit Breaker Test Report", page_icon="📄", layout="wide")


def write_cell(cell, value: str) -> None:
    """Replace a table-cell value while retaining the template cell formatting."""
    value = "" if value is None else str(value)
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(value)
        return

    first = paragraphs[0]
    if first.runs:
        first.runs[0].text = value
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(value)

    for paragraph in paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def replace_paragraph_text(paragraph, value: str) -> None:
    """Replace text while retaining the paragraph's original formatting."""
    value = "" if value is None else str(value)
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def build_report(data: dict) -> bytes:
    doc = Document(TEMPLATE_PATH)
    tables = doc.tables

    # Project information
    project_map = {
        (0, 0, 2): "customer", (0, 0, 4): "order_no",
        (0, 1, 2): "location", (0, 1, 4): "date",
        (0, 2, 2): "pnl_no", (0, 2, 4): "pnl_name",
    }

    # General data
    general_map = {
        (1, 0, 1): "serial_no", (1, 0, 3): "type",
        (1, 1, 1): "mfg_year", (1, 1, 3): "make",
        (1, 2, 1): "rated_voltage", (1, 2, 3): "rated_current",
        (1, 3, 1): "operation_counter", (1, 3, 3): "rated_frequency",
        (1, 4, 1): "control_voltage", (1, 4, 3): "spring_motor_voltage",
        (1, 5, 1): "rated_breaking_capacity", (1, 5, 3): "rated_making_capacity",
    }

    breaker_map = {
        (3, 0, 1): "closing_coil_resistance",
        (3, 0, 3): "spring_motor_resistance",
        (3, 1, 1): "opening_coil_resistance",
    }

    contact_map = {
        (4, 2, 0): "injected_current",
        (4, 2, 1): "contact_r", (4, 2, 2): "contact_y", (4, 2, 3): "contact_b",
        (4, 2, 4): "contact_acceptable",
    }

    timing_map = {
        (5, 1, 1): "close_r", (5, 1, 2): "close_y", (5, 1, 3): "close_b",
        (5, 1, 4): "close_acceptable", (5, 1, 5): "close_control_voltage",
        (5, 2, 1): "open_r", (5, 2, 2): "open_y", (5, 2, 3): "open_b",
        (5, 2, 4): "open_acceptable", (5, 2, 5): "open_control_voltage",
    }

    insulation_map = {
        (6, 2, 1): "rr_before", (6, 2, 2): "rr_after",
        (6, 3, 1): "yy_before", (6, 3, 2): "yy_after",
        (6, 4, 1): "bb_before", (6, 4, 2): "bb_after",
        (6, 2, 4): "ry_before", (6, 2, 5): "ry_after",
        (6, 3, 4): "yb_before", (6, 3, 5): "yb_after",
        (6, 4, 4): "br_before", (6, 4, 5): "br_after",
        (6, 2, 7): "re_before", (6, 2, 8): "re_after",
        (6, 3, 7): "ye_before", (6, 3, 8): "ye_after",
        (6, 4, 7): "be_before", (6, 4, 8): "be_after",
    }

    hv_map = {
        (7, 1, 1): "hv_r", (7, 1, 2): "hv_y", (7, 1, 3): "hv_b",
        (7, 1, 4): "hv_acceptable",
    }

    signature_map = {
        (8, 1, 1): "performed_by_name", (8, 1, 2): "witnessed_by_name",
        (8, 2, 1): "performed_signature_date", (8, 2, 2): "witnessed_signature_date",
    }

    for mapping in (project_map, general_map, breaker_map, contact_map, timing_map, insulation_map, hv_map, signature_map):
        for (table_idx, row_idx, col_idx), key in mapping.items():
            write_cell(tables[table_idx].cell(row_idx, col_idx), data.get(key, ""))

    # Checklist comments: all values are entered manually by the user.
    for row_idx in range(1, 22):
        write_cell(tables[2].cell(row_idx, 2), data.get(f"checklist_{row_idx}", ""))

    # Comments area: write into the blank paragraphs immediately below section 8.
    comments = data.get("comments", "")
    comment_lines = comments.splitlines() or [""]
    start_idx = 24
    available = list(range(start_idx, min(start_idx + 10, len(doc.paragraphs))))
    for i, paragraph_idx in enumerate(available):
        replace_paragraph_text(doc.paragraphs[paragraph_idx], comment_lines[i] if i < len(comment_lines) else "")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


st.title("Circuit Breaker Test Report Generator")
st.caption("Enter values manually and download the completed report in the original Word format.")

with st.form("breaker_report_form"):
    st.subheader("Project Information")
    c1, c2 = st.columns(2)
    with c1:
        customer = st.text_input("Customer")
        location = st.text_input("Location")
        pnl_no = st.text_input("PNL No.")
    with c2:
        order_no = st.text_input("Order No.")
        date = st.text_input("Date")
        pnl_name = st.text_input("PNL Name")

    st.subheader("1. General Data and Information")
    general_labels = [
        ("serial_no", "Serial No."), ("type", "Type"),
        ("mfg_year", "Manufacturing Year"), ("make", "Make"),
        ("rated_voltage", "Rated Voltage"), ("rated_current", "Rated Current"),
        ("operation_counter", "Operation Counter"), ("rated_frequency", "Rated Frequency"),
        ("control_voltage", "Control Voltage"), ("spring_motor_voltage", "Spring Charging Motor Voltage"),
        ("rated_breaking_capacity", "Rated Breaking Capacity"), ("rated_making_capacity", "Rated Making Capacity"),
    ]
    general_values = {}
    cols = st.columns(2)
    for i, (key, label) in enumerate(general_labels):
        with cols[i % 2]:
            general_values[key] = st.text_input(label, key=key)

    st.subheader("2. Checklist — Manual Entry")
    st.info("No predefined options are added. Type each checklist comment/value manually.")
    checklist_descriptions = [
        "Visual Inspection for any physical damage or defects.",
        "Check nameplate information for correctness.",
        "Cleaning of circuit breaker with cleaning agent",
        "Lubrication of operating mechanism & other movable parts",
        "Checking & rectification of operating mechanism",
        "Checking of spring charging Manually & Electrically",
        "Checking of ON-OFF operation Manually & Electrically",
        "Checking & rectification of breaker contact arm assembly & Tulip contact tightness",
        "Greasing of Tulip Contact with Contact Grease",
        "Checking tightness of fasteners, wiring, motor, coils, limit switches and auxiliary contacts",
        "Checking of mechanical interlocks",
        "a. CB can be made ON only in Test & Service position",
        "b. CB cannot be made ON in INTERMEDIATE position",
        "c. Racking operation can be performed in OFF CB",
        "d. Racking operation cannot be performed in ON CB",
        "Checking & Rectification of racking operation",
        "Checking & rectification of Limit Switches",
        "Checking closing & opening operation in test and service position",
        "Checking Breaker 64 Pin plug socket wire & pins",
        "Checking proper working of VCB trolley",
        "Checking shock absorber/damper setting",
    ]
    checklist_values = {}
    for i, description in enumerate(checklist_descriptions, start=1):
        checklist_values[f"checklist_{i}"] = st.text_input(description, key=f"checklist_{i}")

    st.subheader("3. Breaker Data")
    b1, b2, b3 = st.columns(3)
    with b1:
        closing_coil_resistance = st.text_input("Resistance of Closing Coil")
    with b2:
        spring_motor_resistance = st.text_input("Resistance of Spring Charging Motor")
    with b3:
        opening_coil_resistance = st.text_input("Resistance of Opening Coil")

    st.subheader("4. Contact Resistance Test")
    cols = st.columns(5)
    contact_fields = [
        ("injected_current", "Injected Test Current"), ("contact_r", "R"),
        ("contact_y", "Y"), ("contact_b", "B"), ("contact_acceptable", "Acceptable Value"),
    ]
    contact_values = {}
    for col, (key, label) in zip(cols, contact_fields):
        with col:
            contact_values[key] = st.text_input(label, key=key)

    st.subheader("5. Timing Test")
    timing_values = {}
    for operation in ("Close", "Open"):
        st.markdown(f"**{operation}**")
        cols = st.columns(5)
        for col, suffix, label in zip(cols, ["r", "y", "b", "acceptable", "control_voltage"], ["R", "Y", "B", "Acceptable Value", "Control Voltage"]):
            key = f"{operation.lower()}_{suffix}"
            with col:
                timing_values[key] = st.text_input(label, key=key)

    st.subheader("6. Insulation Resistance Value with 5 kV Megger")
    insulation_values = {}
    tests = [
        ("rr", "R – R′"), ("yy", "Y – Y′"), ("bb", "B – B′"),
        ("ry", "R – Y"), ("yb", "Y – B"), ("br", "B – R"),
        ("re", "R – E"), ("ye", "Y – E"), ("be", "B – E"),
    ]
    for key, label in tests:
        c1, c2 = st.columns(2)
        with c1:
            insulation_values[f"{key}_before"] = st.text_input(f"{label} — Before HV", key=f"{key}_before")
        with c2:
            insulation_values[f"{key}_after"] = st.text_input(f"{label} — After HV", key=f"{key}_after")

    st.subheader("7. HV Test")
    cols = st.columns(4)
    hv_values = {}
    for col, key, label in zip(cols, ["hv_r", "hv_y", "hv_b", "hv_acceptable"], ["R Leakage Current", "Y Leakage Current", "B Leakage Current", "Acceptable Value"]):
        with col:
            hv_values[key] = st.text_input(label, key=key)

    st.subheader("8. Comments & Recommendations")
    comments = st.text_area("Comments", height=160)

    st.subheader("Test Personnel")
    c1, c2 = st.columns(2)
    with c1:
        performed_by_name = st.text_input("Test performed by — Name")
        performed_signature_date = st.text_input("Test performed by — Signature with Date")
    with c2:
        witnessed_by_name = st.text_input("Test witnessed — Name")
        witnessed_signature_date = st.text_input("Test witnessed — Signature with Date")

    submitted = st.form_submit_button("Generate Word Document", use_container_width=True)

if submitted:
    data = {
        "customer": customer, "order_no": order_no, "location": location, "date": date,
        "pnl_no": pnl_no, "pnl_name": pnl_name,
        **general_values, **checklist_values,
        "closing_coil_resistance": closing_coil_resistance,
        "spring_motor_resistance": spring_motor_resistance,
        "opening_coil_resistance": opening_coil_resistance,
        **contact_values, **timing_values, **insulation_values, **hv_values,
        "comments": comments,
        "performed_by_name": performed_by_name,
        "performed_signature_date": performed_signature_date,
        "witnessed_by_name": witnessed_by_name,
        "witnessed_signature_date": witnessed_signature_date,
    }
    try:
        report_bytes = build_report(data)
        filename_customer = "_".join(customer.strip().split()) if customer.strip() else "Completed"
        st.success("Word document generated successfully.")
        st.download_button(
            "Download Word Document",
            data=report_bytes,
            file_name=f"{filename_customer}_Circuit_Breaker_Test_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except Exception as exc:
        st.error(f"Could not generate the document: {exc}")
